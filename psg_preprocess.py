# import sys
# import re
# import shutil
# import subprocess
# import argparse
# from pathlib import Path
# from tqdm import tqdm

# # ---------------------------------------------------------------------------
# # KONVERTERING (LibreOffice i stedet for Word)
# # ---------------------------------------------------------------------------

# def _find_libreoffice():
#     """Finder LibreOffice i en VM (Linux eller Windows)."""
#     for name in ("soffice", "libreoffice"):
#         path = shutil.which(name)
#         if path: return path
#     # Standard Windows-stier hvis shutil.which fejler
#     for candidate in [r"C:\Program Files\LibreOffice\program\soffice.exe", 
#                       r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]:
#         if Path(candidate).exists(): return candidate
#     return None

# def convert_doc_to_docx(doc_path, docx_dir):
#     lo_bin = _find_libreoffice()
#     if not lo_bin:
#         raise RuntimeError("LibreOffice blev ikke fundet. Installer det i din VM (sudo apt install libreoffice).")
    
#     # Kør konvertering
#     subprocess.run([
#         lo_bin, "--headless", "--convert-to", "docx", 
#         "--outdir", str(docx_dir), str(doc_path.resolve())
#     ], capture_output=True, check=True)
    
#     return docx_dir / (doc_path.stem + ".docx")

# # ---------------------------------------------------------------------------
# # HØJ-PRÆCISIONS UDTRÆKNING (Fra dit originale script)
# # ---------------------------------------------------------------------------

# _REPLACEMENTS = {
#     "\u00b5": "u", "\u00b1": "+/-", "\u2013": "-", "\u2014": "-",
#     "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u00b0": " grader",
#     "\u2265": ">=", "\u2264": "<=", "\u00d7": "x"
# }

# def normalize_text(text):
#     for char, replacement in _REPLACEMENTS.items():
#         text = text.replace(char, replacement)
#     text = re.sub(r"[^\S\n]", " ", text) # Non-breaking spaces osv.
#     text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f-\x9f]", "", text) # Kontroltegn
#     text = re.sub(r" {2,}", " ", text)
#     return text.strip()

# def is_merge_continuation(tc):
#     """Tjekker om cellen er en del af en flettet celle (skal ignoreres)."""
#     from docx.oxml.ns import qn
#     tcPr = tc.find(qn("w:tcPr"))
#     if tcPr is not None:
#         for tag in ["w:hMerge", "w:vMerge"]:
#             merge = tcPr.find(qn(tag))
#             if merge is not None and merge.get(qn("w:val"), "restart") != "restart":
#                 return True
#     return False

# def extract_table(table):
#     lines = []
#     for row in table.rows:
#         cells = []
#         for cell in row.cells:
#             if not is_merge_continuation(cell._tc):
#                 txt = normalize_text(cell.text)
#                 cells.append(txt if txt else "—")
        
#         # Rens rækken for ydre tomme markører
#         while cells and cells[0] == "—": cells.pop(0)
#         while cells and cells[-1] == "—": cells.pop()
        
#         line = " | ".join(cells)
#         if line and line != "—": lines.append(line)
#     return lines

# def extract_textboxes(doc):
#     from docx.oxml.ns import qn
#     texts = []
#     # Clark-notation for namespaces der ofte bruges til tekstbokse
#     tags = ["{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx",
#             "{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}txbx",
#             "{urn:schemas-microsoft-com:vml}textbox"]
#     for tag in tags:
#         for txbx in doc.element.body.findall(".//" + tag):
#             for p in txbx.findall(".//" + qn("w:p")):
#                 txt = normalize_text("".join(r.text or "" for r in p.findall(".//" + qn("w:t"))))
#                 if txt: texts.append(txt)
#     return texts

# def extract_from_container(container):
#     """Hjælpefunktion til at trække tekst ud af headers/footers rækker og tabeller."""
#     lines = []
#     # Tag alle paragraffer i headeren/footeren
#     for p in container.paragraphs:
#         txt = normalize_text(p.text)
#         if txt:
#             lines.append(txt)
#     # Tag alle tabeller i headeren/footeren
#     for t in container.tables:
#         lines.extend(extract_table(t))
#     return "\n".join(lines)

# def extract_high_fidelity(path):
#     from docx import Document
#     from docx.oxml.ns import qn
#     from docx.text.paragraph import Paragraph
#     from docx.table import Table

#     doc = Document(str(path))
#     output = []

#     # 1. Sidehoved (Rettet her!)
#     header_texts = []
#     for section in doc.sections:
#         # Vi tjekker alle typer headers (normal, første side, lige sider)
#         for h_obj in [section.header, section.first_page_header, section.even_page_header]:
#             if h_obj:
#                 txt = extract_from_container(h_obj)
#                 if txt and txt not in header_texts:
#                     header_texts.append(txt)
    
#     if header_texts:
#         output.append("[SIDEHOVED]\n" + "\n".join(header_texts))

#     # 2. Body (Paragraffer og tabeller i rækkefølge)
#     for child in doc.element.body:
#         if child.tag == qn("w:p"):
#             txt = normalize_text(Paragraph(child, doc).text)
#             if txt: output.append(txt)
#         elif child.tag == qn("w:tbl"):
#             output.extend(extract_table(Table(child, doc)))

#     # 3. Tekstbokse
#     boxes = extract_textboxes(doc)
#     if boxes:
#         output.append("\n[TEKSTBOKSE]\n" + "\n".join(boxes))

#     # 4. Sidefod (Rettet her!)
#     footer_texts = []
#     for section in doc.sections:
#         for f_obj in [section.footer, section.first_page_footer, section.even_page_footer]:
#             if f_obj:
#                 txt = extract_from_container(f_obj)
#                 if txt and txt not in footer_texts:
#                     footer_texts.append(txt)

#     if footer_texts:
#         output.append("\n[SIDEFOD]\n" + "\n".join(footer_texts))

#     return f"=== PSG RAPPORT ===\nFilnavn: {path.name}\n\n" + "\n\n".join(output) + "\n\n=== SLUT ==="

# # ---------------------------------------------------------------------------
# # MAIN PROCESS
# # ---------------------------------------------------------------------------

# def main():
#     parser = argparse.ArgumentParser()
#     parser.add_argument("--mappe", "-m", required=True)
#     parser.add_argument("--liste", "-l", required=True)
#     parser.add_argument("--max-files", type=int)
#     args = parser.parse_args()

#     base_dir = Path(args.mappe)
#     docx_dir = base_dir / "docx"
#     txt_dir = base_dir / "txt"
#     docx_dir.mkdir(parents=True, exist_ok=True)
#     txt_dir.mkdir(parents=True, exist_ok=True)

#     # Indlæs liste
#     stier = [Path(line.split(";")[0].strip()) for line in Path(args.liste).read_text(encoding="utf-8").splitlines() if line.strip()]
#     if args.max_files: stier = stier[:args.max_files]

#     print(f"Starter behandling af {len(stier)} filer...")

#     for sti in tqdm(stier):
#         if not sti.exists(): continue
#         try:
#             # Trin 1: Konvertér (hvis nødvendigt)
#             if sti.suffix.lower() == ".doc":
#                 docx_path = convert_doc_to_docx(sti, docx_dir)
#             else:
#                 # Kopier docx til docx-mappen for at holde det samlet
#                 docx_path = docx_dir / sti.name
#                 shutil.copy2(sti, docx_path)

#             # Trin 2: Ekstrahér med høj præcision
#             text_content = extract_high_fidelity(docx_path)
            
#             # Trin 3: Gem
#             (txt_dir / (sti.stem + ".txt")).write_text(text_content, encoding="utf-8")
            
#         except Exception as e:
#             print(f"\n❌ Fejl ved {sti.name}: {e}")

#     print(f"\n✅ Færdig! Tekstfiler ligger i: {txt_dir}")

# if __name__ == "__main__":
#     main()

import sys
import re
import shutil
import subprocess
import argparse
from pathlib import Path
from tqdm import tqdm

# ---------------------------------------------------------------------------
# OPTIONAL: docx2txt (FAST MODE)
# ---------------------------------------------------------------------------

try:
    import docx2txt
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False


# ---------------------------------------------------------------------------
# KONVERTERING (.doc -> .docx via LibreOffice)
# ---------------------------------------------------------------------------

def _find_libreoffice():
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path

    for candidate in [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
    ]:
        if Path(candidate).exists():
            return candidate

    return None


def convert_doc_to_docx(doc_path, docx_dir):
    lo_bin = _find_libreoffice()
    if not lo_bin:
        raise RuntimeError("LibreOffice ikke fundet.")

    subprocess.run([
        lo_bin,
        "--headless",
        "--convert-to", "docx",
        "--outdir", str(docx_dir),
        str(doc_path.resolve())
    ], capture_output=True, check=True)

    return docx_dir / (doc_path.stem + ".docx")


# ---------------------------------------------------------------------------
# CLEANING
# ---------------------------------------------------------------------------

def deduplicate_lines(text):
    seen = set()
    result = []

    for line in text.split("\n"):
        clean = line.strip()
        if not clean:
            continue

        if clean not in seen:
            seen.add(clean)
            result.append(clean)

    return "\n".join(result)


def normalize_text(text):
    text = re.sub(r"[^\S\n]", " ", text)
    text = re.sub(r"[\x00-\x1f\x7f-\x9f]", "", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# FAST EXTRACTION (docx2txt) 🔥
# ---------------------------------------------------------------------------

def extract_fast(path):
    if not HAS_DOCX2TXT:
        raise RuntimeError("docx2txt mangler: pip install docx2txt")

    text = docx2txt.process(str(path))
    text = normalize_text(text)
    text = deduplicate_lines(text)

    return f"=== PSG RAPPORT ===\nFilnavn: {path.name}\n\n{text}\n\n=== SLUT ==="


# ---------------------------------------------------------------------------
# PRECISE EXTRACTION (din gamle)
# ---------------------------------------------------------------------------

def is_merge_continuation(tc):
    from docx.oxml.ns import qn
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        for tag in ["w:hMerge", "w:vMerge"]:
            merge = tcPr.find(qn(tag))
            if merge is not None and merge.get(qn("w:val"), "restart") != "restart":
                return True
    return False


def extract_table(table):
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            if not is_merge_continuation(cell._tc):
                txt = normalize_text(cell.text)
                cells.append(txt if txt else "—")

        while cells and cells[0] == "—":
            cells.pop(0)
        while cells and cells[-1] == "—":
            cells.pop()

        line = " | ".join(cells)
        if line and line != "—":
            lines.append(line)

    return lines


def extract_textboxes(doc):
    from docx.oxml.ns import qn
    texts = []

    tags = [
        "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx",
        "{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}txbx",
        "{urn:schemas-microsoft-com:vml}textbox"
    ]

    for tag in tags:
        for txbx in doc.element.body.findall(".//" + tag):
            for p in txbx.findall(".//" + qn("w:p")):
                txt = normalize_text("".join(
                    r.text or "" for r in p.findall(".//" + qn("w:t"))
                ))
                if txt:
                    texts.append(txt)

    return texts


def extract_high_fidelity(path):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(str(path))
    output = []

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            txt = normalize_text(Paragraph(child, doc).text)
            if txt:
                output.append(txt)

        elif child.tag == qn("w:tbl"):
            output.extend(extract_table(Table(child, doc)))

    boxes = extract_textboxes(doc)
    if boxes:
        output.append("\n[TEKSTBOKSE]\n" + "\n".join(boxes))

    text = "\n\n".join(output)
    text = deduplicate_lines(text)

    return f"=== PSG RAPPORT ===\nFilnavn: {path.name}\n\n{text}\n\n=== SLUT ==="


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--mappe", "-m", required=True)
    parser.add_argument("--liste", "-l", required=True)
    parser.add_argument("--max-files", type=int)
    parser.add_argument("--mode", choices=["fast", "precise"], default="fast")
    args = parser.parse_args()

    base_dir = Path(args.mappe)
    docx_dir = base_dir / "docx"
    txt_dir = base_dir / "txt"

    docx_dir.mkdir(parents=True, exist_ok=True)
    txt_dir.mkdir(parents=True, exist_ok=True)

    stier = [
        Path(line.split(";")[0].strip())
        for line in Path(args.liste).read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]

    if args.max_files:
        stier = stier[:args.max_files]

    print(f"Starter behandling af {len(stier)} filer...")

    for sti in tqdm(stier):
        if not sti.exists():
            continue

        try:
            if sti.suffix.lower() == ".doc":
                docx_path = convert_doc_to_docx(sti, docx_dir)
            else:
                docx_path = docx_dir / sti.name
                shutil.copy2(sti, docx_path)

            if args.mode == "fast":
                text_content = extract_fast(docx_path)
            else:
                text_content = extract_high_fidelity(docx_path)

            (txt_dir / (sti.stem + ".txt")).write_text(
                text_content, encoding="utf-8"
            )

        except Exception as e:
            print(f"\n❌ Fejl ved {sti.name}: {e}")

    print(f"\n✅ Færdig! Output i: {txt_dir}")


if __name__ == "__main__":
    main()